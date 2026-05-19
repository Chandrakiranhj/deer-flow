#!/usr/bin/env python3
"""Vision Empower premium funder report — vibe-driven DOCX builder.

Same vibe registry as the PPTX/PDF builders. Adds:
  - Cover page with a coloured hero band (table-cell shading)
  - Section heading with eyebrow + accent rule + uppercase tag
  - KPI tiles in a 3-up grid with shaded cells
  - Deliverable progress rendered as a percentage table AND a
    visual bar row (split cells whose widths encode the percentage)
  - Repeating page header (project + funder)
  - Footer with page numbers (Word field codes)
  - Vibe-aware typography (Word substitutes when designer fonts missing)

Usage:
    python build_docx.py --data data.json --output out.docx \\
                         --vibe dark-premium --report-type quarterly \\
                         --period-start 2025-10-01 --period-end 2025-12-31
"""
from __future__ import annotations

import argparse
import json
import sys
import traceback
from datetime import datetime
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from _bootstrap import ensure  # type: ignore

ensure({"docx": "python-docx"})

from docx import Document  # noqa: E402
from docx.enum.table import WD_ALIGN_VERTICAL  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Cm, Mm, Pt, RGBColor, Twips  # noqa: E402

from vibes import Vibe, get_vibe  # noqa: E402


# ── Helpers ───────────────────────────────────────────────────────────────────
def _hex_to_rgb(s: str) -> RGBColor:
    s = s.lstrip("#")
    return RGBColor(int(s[0:2], 16), int(s[2:4], 16), int(s[4:6], 16))


class P:
    """Short-hand palette accessor bound to one vibe — mirrors PPTX/PDF P."""

    def __init__(self, vibe: Vibe):
        p = vibe.palette
        self.bg = _hex_to_rgb(p.bg)
        self.surface = _hex_to_rgb(p.surface)
        self.surface_alt = _hex_to_rgb(p.surface_alt)
        self.border = _hex_to_rgb(p.border)
        self.text_primary = _hex_to_rgb(p.text_primary)
        self.text_secondary = _hex_to_rgb(p.text_secondary)
        self.text_muted = _hex_to_rgb(p.text_muted)
        self.accent = _hex_to_rgb(p.accent)
        self.accent_soft = _hex_to_rgb(p.accent_soft)
        self.danger = _hex_to_rgb(p.danger)
        self.success = _hex_to_rgb(p.success)
        self.hex_bg = p.as_hex("bg")
        self.hex_text_primary = p.as_hex("text_primary")
        self.hex_accent = p.as_hex("accent")
        self.hex_accent_soft = p.as_hex("accent_soft")
        self.hex_surface = p.as_hex("surface")
        self.hex_surface_alt = p.as_hex("surface_alt")
        self.hex_border = p.as_hex("border")


def _resolve_font(name: str) -> str:
    """Word substitutes the named font with an installed family at render
    time if missing — we keep intent (serif vs sans) but rely on the host
    to swap; Fraunces → Georgia, Aptos Display → Calibri, etc."""
    aliases = {
        "Fraunces": "Fraunces",
        "Aptos": "Aptos",
        "Aptos Display": "Aptos Display",
        "Arial Black": "Arial Black",
        "Georgia": "Georgia",
        "Calibri": "Calibri",
    }
    return aliases.get(name, name)


def fmt_money(n) -> str:
    if n is None:
        return "—"
    n = float(n)
    if n >= 1e7:
        return f"₹{n / 1e7:.1f}Cr"
    if n >= 1e5:
        return f"₹{n / 1e5:.1f}L"
    if n >= 1e3:
        return f"₹{n / 1e3:.0f}K"
    return f"₹{int(n):,}"


def fmt_date(d) -> str:
    if not d:
        return "—"
    try:
        dt = datetime.fromisoformat(d.replace("Z", "+00:00"))
        return dt.strftime("%-d %b %Y") if sys.platform != "win32" else dt.strftime("%#d %b %Y")
    except Exception:
        return str(d)


def truncate(s, n: int) -> str:
    if not s:
        return ""
    s = str(s)
    if len(s) <= n:
        return s
    budget = n - 1
    cut = s.rfind(" ", 0, budget)
    if cut < max(8, int(budget * 0.55)):
        cut = budget
    return s[:cut].rstrip(" ,.;:—–-") + "…"


def shade_cell(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def set_cell_borders(cell, hex_border: str, sides=("top", "bottom", "left", "right")) -> None:
    """Set per-side single-line borders for a cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), hex_border)
        tcBorders.append(b)
    tc_pr.append(tcBorders)


def remove_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "bottom", "left", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        tcBorders.append(b)
    tc_pr.append(tcBorders)


def styled_run(p, text: str, *, size=11, bold=False, italic=False,
               color: RGBColor | None = None, font: str = "Aptos"):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color is not None:
        r.font.color.rgb = color
    r.font.name = _resolve_font(font)
    return r


def add_accent_rule(doc, hex_color: str, *, weight=8, space_after_pt=8) -> None:
    """Add a thin horizontal coloured rule below the previous paragraph."""
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(space_after_pt)
    pPr = rule._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(weight))
    bot.set(qn("w:color"), hex_color)
    pBdr.append(bot)
    pPr.append(pBdr)


def add_eyebrow_para(doc, text: str, *, palette: P, vibe: Vibe) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    styled_run(p, text, size=9, bold=True, color=palette.accent,
               font=vibe.typography.body)


def add_section_heading(doc, title: str, eyebrow: str, *,
                        palette: P, vibe: Vibe) -> None:
    """Eyebrow + accent rule + section title — same recipe across PDF/DOCX."""
    add_eyebrow_para(doc, eyebrow.upper(), palette=palette, vibe=vibe)
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(2)
    h.paragraph_format.space_after = Pt(4)
    styled_run(h, title, size=18, bold=True,
               color=palette.text_primary,
               italic=vibe.typography.headline_italic,
               font=vibe.typography.headline)
    add_accent_rule(doc, palette.hex_accent, weight=10, space_after_pt=10)


def add_kpi_grid(doc, values: list[tuple[str, str, str]], *, palette: P,
                 vibe: Vibe) -> None:
    """3-up KPI grid: (label_above, big_value, label_below)."""
    table = doc.add_table(rows=1, cols=len(values))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, (above, val, below) in enumerate(values):
        cell = table.rows[0].cells[i]
        cell.text = ""
        shade_cell(cell, palette.hex_surface_alt)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Above label
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(p1, above.upper(), size=8, bold=True,
                   color=palette.text_muted,
                   font=vibe.typography.body)
        # Big value
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(4)
        styled_run(p2, val, size=26, bold=True,
                   italic=vibe.typography.headline_italic,
                   color=palette.text_primary,
                   font=vibe.typography.headline)
        # Below label
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_before = Pt(2)
        styled_run(p3, below, size=8,
                   color=palette.text_muted,
                   font=vibe.typography.body)
        # Decorative accent rule above each cell.
        tc_pr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side, weight in (("top", 16), ("bottom", 0), ("left", 0), ("right", 0)):
            b = OxmlElement(f"w:{side}")
            if weight > 0:
                b.set(qn("w:val"), "single")
                b.set(qn("w:sz"), str(weight))
                b.set(qn("w:color"), palette.hex_accent)
            else:
                b.set(qn("w:val"), "nil")
            tcBorders.append(b)
        tc_pr.append(tcBorders)


def add_progress_bar_row(table, label: str, pct: int, *, palette: P,
                         vibe: Vibe, label_w_cm=6.5, bar_w_cm=8.0) -> None:
    """Append a 'label | filled-bar | empty-bar | pct%' row.

    The filled/empty split is encoded as two cells whose widths reflect the
    percentage — a low-fi but reliable way to show progress in DOCX without
    embedded charts."""
    row = table.add_row()
    cells = row.cells
    # Adjust widths.
    pct = max(0, min(100, int(pct)))
    fill_w = bar_w_cm * pct / 100
    empty_w = bar_w_cm - fill_w
    cells[0].width = Cm(label_w_cm)
    cells[1].width = Cm(max(0.01, fill_w))
    cells[2].width = Cm(max(0.01, empty_w))
    cells[3].width = Cm(1.6)
    # Label.
    cells[0].text = ""
    styled_run(cells[0].paragraphs[0], label, size=10,
               color=palette.text_secondary, font=vibe.typography.body)
    # Filled portion: shaded accent, no text.
    cells[1].text = ""
    if fill_w > 0:
        shade_cell(cells[1], palette.hex_accent)
    remove_cell_borders(cells[1])
    # Empty track: shaded surface_alt.
    cells[2].text = ""
    if empty_w > 0:
        shade_cell(cells[2], palette.hex_surface_alt)
    remove_cell_borders(cells[2])
    # Percentage label.
    cells[3].text = ""
    p = cells[3].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    styled_run(p, f"{pct}%", size=10, bold=True,
               color=palette.accent, font=vibe.typography.body)


def add_page_field(p, instr: str) -> None:
    """Add a Word field code (PAGE, NUMPAGES, etc.) to a paragraph."""
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" {instr} "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = p.add_run()
    r._r.append(fld_begin)
    r._r.append(instr_text)
    r._r.append(fld_end)


def configure_header_footer(doc, project_name: str, funder: str, *,
                            palette: P, vibe: Vibe) -> None:
    """Set the repeating page header and the page-numbered footer.
    Different-first-page is on so the cover gets no header/footer."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # Header
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    styled_run(header_p, f"{project_name}  ", size=8, bold=True,
               color=palette.text_secondary, font=vibe.typography.body)
    styled_run(header_p, f"·  {funder}", size=8,
               color=palette.text_muted, font=vibe.typography.body)
    # Thin accent rule under the header.
    pPr = header_p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "8")
    bot.set(qn("w:color"), palette.hex_accent)
    pBdr.append(bot)
    pPr.append(pBdr)

    # Footer
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    styled_run(footer_p, "Vision Empower Trust  ·  ", size=8,
               italic=True, color=palette.text_muted,
               font=vibe.typography.body)
    styled_run(footer_p, "www.visionempower.in", size=8,
               italic=True, color=palette.text_muted,
               font=vibe.typography.body)
    # Tab + "Page X" right-aligned via a manual tab stop.
    from docx.shared import Emu
    tab_stops = footer_p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(16))
    styled_run(footer_p, "\tPage ", size=8,
               color=palette.text_muted, font=vibe.typography.body)
    add_page_field(footer_p, "PAGE")
    styled_run(footer_p, " of ", size=8,
               color=palette.text_muted, font=vibe.typography.body)
    add_page_field(footer_p, "NUMPAGES")


def add_cover_band(doc, *, project_name: str, funder: str, period: str,
                   grant_amount, report_type_label: str, palette: P,
                   vibe: Vibe) -> None:
    """Cover hero — a single-row table whose cell is shaded with the vibe's
    primary text colour and contains the project name + meta."""
    band = doc.add_table(rows=1, cols=1)
    band.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = band.rows[0].cells[0]
    cell.width = Cm(17)
    cell.text = ""
    # Top accent rule on the band.
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, sz, col in (
        ("top", 24, palette.hex_accent),
        ("bottom", 0, "auto"),
        ("left", 0, "auto"),
        ("right", 0, "auto"),
    ):
        b = OxmlElement(f"w:{side}")
        if sz > 0:
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), str(sz))
            b.set(qn("w:color"), col)
        else:
            b.set(qn("w:val"), "nil")
        tcBorders.append(b)
    tc_pr.append(tcBorders)
    # Band fill = primary text colour for high-contrast headline against it.
    shade_cell(cell, palette.hex_text_primary)

    # Eyebrow
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(20)
    p1.paragraph_format.left_indent = Cm(0.6)
    styled_run(p1, "VISION EMPOWER", size=10, bold=True,
               color=palette.accent_soft, font=vibe.typography.body)

    p2 = cell.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.6)
    p2.paragraph_format.space_after = Pt(6)
    styled_run(p2, report_type_label, size=9, bold=True,
               color=palette.accent_soft, font=vibe.typography.body)

    # Title
    p3 = cell.add_paragraph()
    p3.paragraph_format.left_indent = Cm(0.6)
    p3.paragraph_format.space_before = Pt(8)
    p3.paragraph_format.space_after = Pt(12)
    styled_run(p3, project_name, size=32, bold=True,
               italic=vibe.typography.headline_italic,
               color=_hex_to_rgb("#FFFFFF"),
               font=vibe.typography.headline)

    # Metadata strip
    p4 = cell.add_paragraph()
    p4.paragraph_format.left_indent = Cm(0.6)
    p4.paragraph_format.space_after = Pt(20)
    styled_run(p4, f"Funder  ", size=9, bold=True,
               color=palette.accent_soft, font=vibe.typography.body)
    styled_run(p4, f"{funder}    ", size=10,
               color=_hex_to_rgb("#FFFFFF"), font=vibe.typography.body)
    styled_run(p4, "Period  ", size=9, bold=True,
               color=palette.accent_soft, font=vibe.typography.body)
    styled_run(p4, f"{period}    ", size=10,
               color=_hex_to_rgb("#FFFFFF"), font=vibe.typography.body)
    styled_run(p4, "Grant  ", size=9, bold=True,
               color=palette.accent_soft, font=vibe.typography.body)
    styled_run(p4, fmt_money(grant_amount), size=10,
               color=_hex_to_rgb("#FFFFFF"), font=vibe.typography.body)


# ── Main ──────────────────────────────────────────────────────────────────────
def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--data", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--report-type", default="quarterly")
    parser.add_argument("--period-start", default="")
    parser.add_argument("--period-end", default="")
    parser.add_argument("--draft", default="")
    parser.add_argument("--vibe", default="editorial-serif")
    args = parser.parse_args(argv)

    data = json.loads(Path(args.data).read_text(encoding="utf-8"))
    project = data.get("project") or data
    if "budgets" not in project and project.get("budgetCategories"):
        project["budgets"] = project["budgetCategories"]
    if "activities" not in project and project.get("recentActivities"):
        project["activities"] = project["recentActivities"]
    if "deliverablesTotal" not in project or "deliverablesDone" not in project:
        delivs = project.get("deliverables") or []
        project.setdefault("deliverablesTotal", len(delivs))
        project.setdefault(
            "deliverablesDone",
            sum(1 for d in delivs if (d.get("achieved") or 0) >= (d.get("target") or 0) > 0),
        )
    if "approvedBudget" not in project or "spentBudget" not in project:
        cats = project.get("budgetCategories") or project.get("budgets") or []
        project.setdefault(
            "approvedBudget",
            sum((c.get("approvedAmount") or c.get("approved") or 0) for c in cats),
        )
        project.setdefault(
            "spentBudget",
            sum((c.get("spentAmount") or c.get("spent") or 0) for c in cats),
        )
    period_start = args.period_start or data.get("periodStart") or project.get("startDate") or ""
    period_end = args.period_end or data.get("periodEnd") or project.get("endDate") or ""
    draft = args.draft or data.get("draft") or ""

    vibe = get_vibe(args.vibe)
    palette = P(vibe)
    BODY_FONT = vibe.typography.body
    HEAD_FONT = vibe.typography.headline

    # ── Section gating mirrors PPTX/PDF manifest ─────────────────────────
    narrative = project.get("narrative") if isinstance(project.get("narrative"), dict) else {}
    summary_text = ((narrative.get("overview") or "").strip() if narrative else "") or project.get("summary") or ""
    deliverables = project.get("deliverables") or []
    activities = project.get("activities") or []
    testimonials = list(project.get("testimonials") or [])
    if not testimonials:
        for a in activities:
            if a.get("testimonial"):
                testimonials.append({
                    "content": a["testimonial"],
                    "author": a.get("testimonialBy") or a.get("title", ""),
                    "role": a.get("state") or "",
                })

    quote_threshold = 40 if vibe.layout.prefers_quotes else 100
    overview_min = 80 if vibe.layout.prefers_prose else 180
    way_forward_min_draft = 300 if vibe.layout.prefers_prose else 600

    has_overview = bool(summary_text and len(summary_text) >= overview_min) or bool(
        draft and len(draft) >= max(200, overview_min * 2)
    )
    has_deliverables = len(deliverables) >= 1
    has_activities = len(activities) >= 1
    has_stories = any(
        (t.get("content") or "").strip() and len((t.get("content") or "").strip()) >= quote_threshold
        for t in testimonials
    )
    approved = project.get("approvedBudget") or 0
    spent = project.get("spentBudget") or 0
    has_financials = approved > 0 or (vibe.layout.prefers_charts and (spent > 0 or bool(project.get("budgets"))))
    has_narrative_forward = (
        (narrative.get("challenges") or "").strip()
        or (narrative.get("way_forward") or "").strip()
    ) if narrative else False
    has_way_forward = has_narrative_forward or bool(draft and len(draft) >= way_forward_min_draft)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    project_name = project.get("name", "")
    funder = project.get("funderName") or "Vision Empower"
    is_quarterly = args.report_type == "quarterly"
    period = (
        f"{fmt_date(period_start)} – {fmt_date(period_end)}" if is_quarterly
        else f"{fmt_date(project.get('startDate'))} – {fmt_date(project.get('endDate'))}"
    )
    report_type_label = "QUARTERLY FUNDER REPORT" if is_quarterly else "FULL PROJECT REPORT"

    # ── Header / footer (first page exempted by `different_first_page`) ──
    configure_header_footer(doc, project_name, funder,
                            palette=palette, vibe=vibe)

    # ── Cover band ───────────────────────────────────────────────────────
    add_cover_band(doc, project_name=project_name, funder=funder,
                   period=period, grant_amount=project.get("grantAmount"),
                   report_type_label=report_type_label,
                   palette=palette, vibe=vibe)

    # Cover sub-content (states, summary preview).
    if project.get("states"):
        st = doc.add_paragraph()
        st.paragraph_format.space_before = Pt(20)
        styled_run(st, "STATES   ", size=9, bold=True,
                   color=palette.accent, font=BODY_FONT)
        styled_run(st, ", ".join(project["states"]), size=11,
                   color=palette.text_secondary, font=BODY_FONT)

    if summary_text:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(6)
        sp.paragraph_format.space_after = Pt(0)
        styled_run(sp, summary_text[:300], size=11, italic=True,
                   color=palette.text_secondary, font=BODY_FONT)

    # Page break onto first content page (header/footer kick in here).
    doc.add_page_break()

    # ── Executive summary (gated) ────────────────────────────────────────
    if has_overview:
        add_section_heading(doc, "Executive Summary", "Overview",
                            palette=palette, vibe=vibe)
        summary = summary_text
        if not summary and draft:
            lines = [l.strip() for l in draft.split("\n") if len(l.strip()) > 60]
            summary = " ".join(lines[:4])[:600]
        if summary:
            p = doc.add_paragraph()
            styled_run(p, summary, size=11, color=palette.text_secondary,
                       font=BODY_FONT)

    # ── Impact at a glance ──────────────────────────────────────────────
    teachers = sum(a.get("teachersReached") or 0 for a in activities)
    students = sum(a.get("studentsReached") or 0 for a in activities)
    schools = sum(a.get("schoolsReached") or 0 for a in activities)
    if teachers or students or schools:
        add_section_heading(doc, "Impact at a Glance", "This Period",
                            palette=palette, vibe=vibe)
        add_kpi_grid(doc, [
            ("Teachers", f"{teachers:,}", "reached this period"),
            ("Students", f"{students:,}", "reached this period"),
            ("Schools",  f"{schools:,}",  "reached this period"),
        ], palette=palette, vibe=vibe)

    # ── Deliverables: table + visual bars ────────────────────────────────
    if has_deliverables:
        add_section_heading(doc, "Deliverable Progress",
                            "What the Grant Committed To",
                            palette=palette, vibe=vibe)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, label in enumerate(["Deliverable", "Target", "Achieved", "% Complete"]):
            hdr[i].text = ""
            styled_run(hdr[i].paragraphs[0], label, size=10, bold=True,
                       color=_hex_to_rgb("#FFFFFF"), font=BODY_FONT)
            shade_cell(hdr[i], palette.hex_text_primary)
        for d in deliverables:
            row = table.add_row().cells
            unit = d.get("unit") or ""
            row[0].text = d.get("title", "")
            row[1].text = f"{d.get('target') or '—'}{(' ' + unit) if (d.get('target') and unit) else ''}"
            row[2].text = f"{d.get('achieved') or '—'}{(' ' + unit) if (d.get('achieved') and unit) else ''}"
            target = d.get("target") or 0
            achieved = d.get("achieved") or 0
            pct = round(achieved / target * 100) if target else (100 if d.get("status") == "completed" else 0)
            row[3].text = f"{pct}%"
            for cell in row:
                for para in cell.paragraphs:
                    for r in para.runs:
                        r.font.name = _resolve_font(BODY_FONT)
                        r.font.color.rgb = palette.text_secondary
                        r.font.size = Pt(10.5)

        # Visual progress bars beneath the table.
        bars_intro = doc.add_paragraph()
        bars_intro.paragraph_format.space_before = Pt(10)
        styled_run(bars_intro, "VISUAL PROGRESS", size=8, bold=True,
                   color=palette.text_muted, font=BODY_FONT)
        bar_table = doc.add_table(rows=0, cols=4)
        for d in deliverables[:8]:
            target = d.get("target") or 0
            achieved = d.get("achieved") or 0
            pct = round(achieved / target * 100) if target else (100 if d.get("status") == "completed" else 0)
            add_progress_bar_row(bar_table, d.get("title", "")[:48], pct,
                                 palette=palette, vibe=vibe)

    # ── Field activities ────────────────────────────────────────────────
    if has_activities:
        add_section_heading(doc, "Field Activities", "On-Ground Work",
                            palette=palette, vibe=vibe)
        for act in activities[:20]:
            p = doc.add_paragraph(style="List Bullet")
            date_str = fmt_date(act.get("activityDate")) if act.get("activityDate") else ""
            head = f"{date_str} — {act.get('title', '')}" if date_str else act.get("title", "")
            styled_run(p, head + ". ", size=11, bold=True,
                       color=palette.text_primary, font=BODY_FONT)
            reaches = []
            if act.get("teachersReached"): reaches.append(f"{act['teachersReached']} teachers")
            if act.get("studentsReached"): reaches.append(f"{act['studentsReached']} students")
            if act.get("schoolsReached"):  reaches.append(f"{act['schoolsReached']} schools")
            tail_bits = []
            if act.get("state"): tail_bits.append(act["state"])
            if reaches: tail_bits.append(", ".join(reaches))
            if act.get("notes"): tail_bits.append(act["notes"])
            if tail_bits:
                styled_run(p, " · ".join(tail_bits), size=11,
                           color=palette.text_secondary, font=BODY_FONT)

    # ── Financials ──────────────────────────────────────────────────────
    if has_financials:
        add_section_heading(doc, "Budget & Utilisation",
                            "Where the Grant Is Going",
                            palette=palette, vibe=vibe)
        util_pct = round(spent / approved * 100) if approved else 0
        # Top KPI strip for finance.
        add_kpi_grid(doc, [
            ("Approved", fmt_money(approved), "grant committed"),
            ("Spent",    fmt_money(spent),    "this period"),
            ("Utilised", f"{util_pct}%",      "of grant"),
        ], palette=palette, vibe=vibe)

        budgets = project.get("budgets") or []
        if budgets:
            bt = doc.add_table(rows=1, cols=4)
            bt.style = "Light Grid Accent 1"
            for i, h in enumerate(["Category", "Approved", "Spent", "Utilisation"]):
                cell = bt.rows[0].cells[i]
                cell.text = ""
                styled_run(cell.paragraphs[0], h, size=10, bold=True,
                           color=_hex_to_rgb("#FFFFFF"), font=BODY_FONT)
                shade_cell(cell, palette.hex_text_primary)
            for b in budgets:
                row = bt.add_row().cells
                row[0].text = b.get("name", "")
                row[1].text = fmt_money(b.get("approvedAmount"))
                row[2].text = fmt_money(b.get("spentAmount"))
                ap = b.get("approvedAmount") or 0
                sp = b.get("spentAmount") or 0
                row[3].text = f"{round(sp / ap * 100) if ap else 0}%"
                for cell in row:
                    for para in cell.paragraphs:
                        for r in para.runs:
                            r.font.name = _resolve_font(BODY_FONT)
                            r.font.color.rgb = palette.text_secondary
                            r.font.size = Pt(10.5)

    # ── Stories (gated) ─────────────────────────────────────────────────
    if has_stories:
        add_section_heading(doc, "Stories from the Field",
                            "In Their Own Words", palette=palette, vibe=vibe)
        strong = [t for t in testimonials
                  if (t.get("content") or "").strip()
                  and len((t.get("content") or "").strip()) >= quote_threshold]
        for t in strong[:3]:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            styled_run(p, f'"{t.get("content", "")}"', size=12, italic=True,
                       color=palette.text_primary, font=HEAD_FONT)
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            role = t.get("role") or ""
            styled_run(p2, f"— {t.get('author', '')}{', ' + role if role else ''}",
                       size=10, color=palette.text_muted, font=BODY_FONT)

    # ── Achievements (if structured narrative provided) ─────────────────
    if narrative and (narrative.get("achievements") or "").strip():
        add_section_heading(doc, "Key Achievements",
                            "What We're Proud Of",
                            palette=palette, vibe=vibe)
        p = doc.add_paragraph()
        styled_run(p, narrative["achievements"].strip()[:1200], size=11,
                   color=palette.text_secondary, font=BODY_FONT)

    # ── Way forward (gated; no filler) ──────────────────────────────────
    if has_way_forward:
        add_section_heading(doc, "Challenges & Next Steps",
                            "Looking Ahead", palette=palette, vibe=vibe)
        challenges = (narrative.get("challenges") or "").strip() if narrative else ""
        next_steps = (narrative.get("way_forward") or "").strip() if narrative else ""

        if (not challenges or not next_steps) and draft:
            import re as _re
            def extract(text, start_re, stop_re):
                parts = _re.split(start_re, text, flags=_re.IGNORECASE, maxsplit=1)
                if len(parts) < 2: return ""
                content = parts[1]
                m = _re.search(stop_re, content, flags=_re.IGNORECASE)
                if m: content = content[: m.start()]
                return _re.sub(r"\s+", " ", _re.sub(r"^[\s:–\-]+", "", content)).strip()[:600]
            if not challenges:
                challenges = extract(draft, r"challenges?|barriers?|constraints?",
                                     r"next\s+steps?|way\s+forward|financial")
            if not next_steps:
                next_steps = extract(draft, r"next\s+steps?|way\s+forward|upcoming\s+activities",
                                     r"evidence|conclusion|thank|challenges?")

        if challenges:
            p = doc.add_paragraph()
            styled_run(p, "Challenges. ", size=11, bold=True,
                       color=palette.text_primary, font=BODY_FONT)
            styled_run(p, challenges, size=11, color=palette.text_secondary,
                       font=BODY_FONT)
        if next_steps:
            p = doc.add_paragraph()
            styled_run(p, "Next Steps. ", size=11, bold=True,
                       color=palette.text_primary, font=BODY_FONT)
            styled_run(p, next_steps, size=11, color=palette.text_secondary,
                       font=BODY_FONT)

    out = Path(args.output)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    size = out.stat().st_size
    if size < 6000:
        raise RuntimeError(f"Generated DOCX is suspiciously small ({size} bytes)")
    print(f"WROTE: {out} ({size:,} bytes, vibe={vibe.key})")
    return 0


if __name__ == "__main__":
    try:
        sys.exit(main())
    except Exception:
        traceback.print_exc()
        sys.exit(1)
